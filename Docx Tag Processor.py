import os
import re
from collections import OrderedDict
from docx import Document
import tkinter as tk
from tkinter import ttk, filedialog, messagebox

def scan_folder(folder_path):
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('.docx')]
    tag_order = OrderedDict()
    for file in docx_files:
        doc = Document(os.path.join(folder_path, file))
        for paragraph in doc.paragraphs:
            tags = re.findall(r'\{(.*?)\}', paragraph.text)
            for tag in tags:
                if tag not in tag_order:
                    tag_order[tag] = None
    return list(tag_order.keys()), docx_files

def replace_tags(doc, replacements):
    for paragraph in doc.paragraphs:
        for tag, value in replacements.items():
            if f'{{{tag}}}' in paragraph.text:
                paragraph.text = paragraph.text.replace(f'{{{tag}}}', value)
    return doc

def process_files(input_folder, output_folder, replacements):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    docx_files = [f for f in os.listdir(input_folder) if f.endswith('.docx')]
    for file in docx_files:
        doc = Document(os.path.join(input_folder, file))
        processed_doc = replace_tags(doc, replacements)
        processed_doc.save(os.path.join(output_folder, file))

def open_folder():
    folder_path = filedialog.askdirectory()
    if folder_path:
        input_folder = os.path.join(folder_path, "Template Files")
        if os.path.exists(input_folder):
            tags, files = scan_folder(input_folder)
            create_input_fields(tags, folder_path)
        else:
            messagebox.showerror("Error", "Template Files folder not found in the selected directory.")

def create_input_fields(tags, folder_path):
    fields = {}

    frame = ttk.Frame(root, padding="10 10 10 10")
    frame.pack(fill=tk.BOTH, expand=True)
    
    for tag in tags:
        tag_frame = ttk.Frame(frame)
        tag_frame.pack(fill=tk.X, pady=5)
        
        label = ttk.Label(tag_frame, text=f'{tag}:')
        label.pack(side=tk.LEFT, padx=5)
        
        entry = ttk.Entry(tag_frame)
        entry.pack(side=tk.RIGHT, padx=5, fill=tk.X, expand=True)
        
        fields[tag] = entry

    def on_submit():
        replacements = {tag: entry.get() for tag, entry in fields.items()}
        output_folder = os.path.join(folder_path, "output")
        input_folder = os.path.join(folder_path, "Template Files")
        process_files(input_folder, output_folder, replacements)
        messagebox.showinfo("Success", "Files processed successfully!")

    submit_button = ttk.Button(root, text="Process Files", command=on_submit)
    submit_button.pack(pady=10)

root = tk.Tk()
root.title("Docx Tag Processor")
root.geometry("400x300")

open_button = ttk.Button(root, text="Open Folder", command=open_folder)
open_button.pack(pady=10)

root.mainloop()
